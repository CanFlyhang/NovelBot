from datetime import date, datetime
from typing import List
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
from sqlalchemy.orm import Session

from .config import settings
from .db import get_db
from .models import Chapter, CreationLog, Novel, NovelStatus, ChapterStatus
from .scheduler import scheduler
from .schemas import (
    ConfigUpdate,
    ControlCommand,
    ControlState,
    DashboardSummary,
    Novel as NovelSchema,
    NovelCreate,
    Chapter as ChapterSchema,
    CreationLog as CreationLogSchema,
)
from .services.novel_service import (
    generate_next_chapter_for_novel,
    get_dashboard_summary,
)


router = APIRouter(prefix="/api")


@router.post("/novels", response_model=NovelSchema)
def create_novel(
    novel_in: NovelCreate,
    db: Session = Depends(get_db),
) -> NovelSchema:
    """
    创建一条新的小说规划记录，用于后续自动生成章节。
    """

    planned_date = novel_in.planned_date or date.today()

    novel = Novel(
        title=novel_in.title,
        genre=novel_in.genre,
        description=novel_in.description,
        target_chapter_count=novel_in.target_chapter_count,
        status=NovelStatus.PLANNED,
        planned_date=planned_date,
    )
    db.add(novel)
    db.flush()

    for idx in range(1, novel_in.target_chapter_count + 1):
        chapter = Chapter(
            novel_id=novel.id,
            index=idx,
            title=f"第{idx}章",
            status=ChapterStatus.PLANNED,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
        )
        db.add(chapter)

    db.commit()
    db.refresh(novel)
    return novel


@router.get("/dashboard", response_model=DashboardSummary)
def get_dashboard(db: Session = Depends(get_db)) -> DashboardSummary:
    """
    获取仪表盘所需的创作进度与统计数据。
    """

    return get_dashboard_summary(db)


@router.get("/novels", response_model=List[NovelSchema])
def list_novels(db: Session = Depends(get_db)) -> List[NovelSchema]:
    """
    获取最近的小说列表及其基本信息。
    """

    novels: List[Novel] = (
        db.query(Novel)
        .order_by(Novel.created_at.desc())
        .limit(100)
        .all()
    )
    return novels


@router.delete("/novels/{novel_id}", response_model=dict)
def delete_novel(
    novel_id: int,
    db: Session = Depends(get_db),
) -> dict:
    """
    删除指定小说及其关联的章节和日志。
    """

    novel: Novel | None = db.query(Novel).get(novel_id)
    if novel is None:
        raise HTTPException(status_code=404, detail="小说不存在")

    db.delete(novel)
    db.commit()
    return {"success": True}


@router.get(
    "/novels/{novel_id}/chapters",
    response_model=List[ChapterSchema],
)
def list_chapters_for_novel(
    novel_id: int,
    db: Session = Depends(get_db),
) -> List[ChapterSchema]:
    """
    获取指定小说已生成的章节列表（按章节顺序）。
    """

    chapters: List[Chapter] = (
        db.query(Chapter)
        .filter(Chapter.novel_id == novel_id, Chapter.content.isnot(None))
        .order_by(Chapter.index.asc())
        .all()
    )
    return chapters


@router.post("/novels/{novel_id}/generate")
def manual_generate_chapter(
    novel_id: int,
    db: Session = Depends(get_db),
) -> dict:
    """
    手动为指定小说生成下一章内容。
    """

    ok = generate_next_chapter_for_novel(db, novel_id)
    if not ok:
        raise HTTPException(status_code=400, detail="无法生成新的章节")
    return {"success": True}


@router.get("/logs", response_model=List[CreationLogSchema])
def list_logs(
    limit: int = 200,
    db: Session = Depends(get_db),
) -> List[CreationLogSchema]:
    """
    获取最新的创作过程日志列表。
    """

    limit = max(1, min(limit, 500))
    logs: List[CreationLog] = (
        db.query(CreationLog)
        .order_by(CreationLog.created_at.desc())
        .limit(limit)
        .all()
    )
    return logs


@router.get("/chapters/{chapter_id}", response_model=ChapterSchema)
def get_chapter(
    chapter_id: int,
    db: Session = Depends(get_db),
) -> ChapterSchema:
    """
    根据章节 ID 查询章节详情及正文内容。
    """

    chapter: Chapter | None = db.query(Chapter).get(chapter_id)
    if chapter is None:
        raise HTTPException(status_code=404, detail="章节不存在")
    return chapter


@router.get("/novels/{novel_id}/export-docx")
def export_novel_docx(
    novel_id: int,
    db: Session = Depends(get_db),
) -> StreamingResponse:
    """
    导出指定小说的所有已生成章节为 Word 文档。
    """

    novel: Novel | None = db.query(Novel).get(novel_id)
    if novel is None:
        raise HTTPException(status_code=404, detail="小说不存在")

    chapters: List[Chapter] = (
        db.query(Chapter)
        .filter(Chapter.novel_id == novel_id, Chapter.content.isnot(None))
        .order_by(Chapter.index.asc())
        .all()
    )
    if not chapters:
        raise HTTPException(status_code=400, detail="该小说尚无可导出的章节")

    document = Document()
    document.add_heading(novel.title, level=1)
    if novel.description:
        document.add_paragraph(novel.description)

    for ch in chapters:
        document.add_heading(f"第{ch.index}章 {ch.title}", level=2)
        if ch.outline:
            document.add_paragraph(f"本章小结：{ch.outline}")
        if ch.content:
            for para in ch.content.split("\n\n"):
                text = para.strip()
                if not text:
                    continue
                document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    safe_title = "".join(
        c if c.isalnum() else "_" for c in (novel.title or "novel")
    )
    filename_ascii = (safe_title or "novel").encode(
        "ascii", "ignore"
    ).decode("ascii")
    encoded_title = quote((novel.title or "novel").replace(" ", "_"))

    content_disposition = (
        f'attachment; filename="{filename_ascii}.docx"; '
        f"filename*=UTF-8''{encoded_title}.docx"
    )

    return StreamingResponse(
        buffer,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": content_disposition,
        },
    )


@router.get("/novels/{novel_id}/latest-chapter", response_model=ChapterSchema)
def get_latest_chapter_for_novel(
    novel_id: int,
    db: Session = Depends(get_db),
) -> ChapterSchema:
    """
    查询指定小说最近一次已生成的章节详情。
    """

    chapter: Chapter | None = (
        db.query(Chapter)
        .filter(Chapter.novel_id == novel_id, Chapter.content.isnot(None))
        .order_by(Chapter.index.desc())
        .first()
    )
    if chapter is None:
        raise HTTPException(status_code=404, detail="该小说暂无已生成章节")
    return chapter


@router.post("/control", response_model=ControlState)
def control_scheduler(cmd: ControlCommand) -> ControlState:
    """
    通过 API 控制调度器的启动、暂停与停止。
    """

    action = cmd.action.lower()
    if action == "start":
        scheduler.start()
    elif action == "pause":
        scheduler.pause()
    elif action == "resume":
        scheduler.resume()
    elif action == "stop":
        scheduler.stop()
    else:
        raise HTTPException(status_code=400, detail="未知控制命令")

    return ControlState(
        is_running=scheduler.is_running(),
        is_paused=scheduler.is_paused(),
        last_heartbeat=scheduler.last_heartbeat(),
    )


@router.get("/control/state", response_model=ControlState)
def get_control_state() -> ControlState:
    """
    查询当前调度器运行状态。
    """

    return ControlState(
        is_running=scheduler.is_running(),
        is_paused=scheduler.is_paused(),
        last_heartbeat=scheduler.last_heartbeat(),
    )


@router.post("/config", response_model=dict)
def update_config(config: ConfigUpdate) -> dict:
    """
    动态更新部分运行参数，例如每日产量目标等。
    """

    if config.daily_target_novels is not None:
        settings.daily_target_novels = config.daily_target_novels
    if config.default_chapters_per_novel is not None:
        settings.default_chapters_per_novel = config.default_chapters_per_novel
    if config.max_concurrent_api_requests is not None:
        settings.max_concurrent_api_requests = (
            config.max_concurrent_api_requests
        )
    if config.max_requests_per_minute is not None:
        settings.max_requests_per_minute = config.max_requests_per_minute
    if config.preferred_genres is not None:
        settings.preferred_genres = config.preferred_genres

    return {"success": True, "updated_at": datetime.utcnow().isoformat()}
